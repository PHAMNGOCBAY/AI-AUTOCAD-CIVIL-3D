#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DAM BE TONG COT THEP — TCVN 5574:2018
b=300 x h=800mm | L=6000mm | 4phi18 (CB400-V) | phi10@150 (CB240-T) | B25
Mu = 500 kN.m
"""
import sys, io, math, os
from datetime import datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT_DIR = r"G:\My Drive\AI-AUTOCAD CIVIL 3D\projects\DamBeTong"
os.makedirs(OUT_DIR, exist_ok=True)
NOW = datetime.now().strftime("%Y-%m-%d %H:%M")

# ═══════════════════════════════════════════════════════════════
# 1. THONG SO DAU VAO
# ═══════════════════════════════════════════════════════════════
b       = 300.0     # Chieu rong dam [mm]  (gia thiet)
h       = 800.0     # Chieu cao dam [mm]
L       = 6000.0    # Chieu dai dam [mm]
a_cv    = 25.0      # Lop bao ve be tong [mm]
Mu      = 500.0     # Mo men tinh toan [kN.m]

# Cot thep chu duoi — CB400-V
n_main  = 4
d_main  = 18.0      # [mm]
Rs      = 350.0     # [MPa] TCVN 5574:2018 Bang 8
Rsc     = 350.0
Es      = 200000.0
xi_R    = 0.531     # TCVN 5574:2018 Bang 9 (CB400-V voi B25)

# Cot cau tao tren — 2phi12
n_top   = 2
d_top   = 12.0

# Cot dai — CB240-T
d_stir  = 10.0
n_legs  = 2         # 2 nhanh
s_stir  = 150.0     # [mm]
Rsw     = 175.0     # [MPa]

# Be tong B25 — TCVN 5574:2018 Bang 6
Rb      = 14.5      # Cuong do chiu nen [MPa]
Rbt     = 1.05      # Cuong do chiu keo [MPa]
Eb      = 30000.0   # Modul dan hoi [MPa]
nu      = 0.2

# ═══════════════════════════════════════════════════════════════
# 2. TINH TOAN HINH HOC
# ═══════════════════════════════════════════════════════════════
As_main = n_main * math.pi / 4 * d_main**2       # [mm²]
As_top  = n_top  * math.pi / 4 * d_top**2        # [mm²]
As_stir = n_legs * math.pi / 4 * d_stir**2       # [mm²] tren 1 cat ngang
a_s     = a_cv + d_stir + d_main / 2             # khoang cach TTT -> mep duoi [mm]
h0      = h - a_s                                 # chieu cao lam viec [mm]
A_sect  = b * h                                   # [mm²]

# ═══════════════════════════════════════════════════════════════
# 3. KIEM TOAN UON — TCVN 5574:2018 muc 8.1.2
# ═══════════════════════════════════════════════════════════════
x_act   = Rs * As_main / (Rb * b)                # [mm]
xi_act  = x_act / h0
ok_xi   = xi_act <= xi_R

MRd_Nmm = Rb * b * x_act * (h0 - x_act / 2)     # [N.mm]
MRd     = MRd_Nmm / 1e6                           # [kN.m]
ratio_M = Mu / MRd
ok_bend = ratio_M <= 1.0

# As yeu cau de chiu Mu
phi_m   = Mu * 1e6 / (Rb * b * h0**2)
phi_m   = min(phi_m, 0.499)
xi_req  = 1.0 - math.sqrt(max(0.0, 1.0 - 2.0 * phi_m))
As_req  = Rb * b * xi_req * h0 / Rs              # [mm²]

# Ham luong thep
As_min  = 0.0013 * b * h0                        # TCVN 5574:2018
mu_pct  = As_main / (b * h0) * 100               # [%]
ok_Asmin = As_main >= As_min

# Goi y so luong thep thay the neu khong dat
if not ok_bend:
    # Tim n*phi26
    As_need = As_req
    for d_try in [20, 22, 25, 28, 32]:
        A1 = math.pi / 4 * d_try**2
        n_try = math.ceil(As_need / A1)
        n_try = max(n_try, 4)
        As_try = n_try * A1
        x_try = Rs * As_try / (Rb * b)
        MRd_try = Rb * b * x_try * (h0 - x_try/2) / 1e6
        if MRd_try >= Mu:
            goi_y = f"{n_try}phi{d_try} (As={As_try:.0f} mm2, MRd={MRd_try:.1f} kN.m)"
            break
    else:
        goi_y = f"Can tang kich thuoc mat cat hoac tang b"
else:
    goi_y = "Khong can thay the"

# ═══════════════════════════════════════════════════════════════
# 4. KHOI LUONG
# ═══════════════════════════════════════════════════════════════
# Be tong
V_bt_m3    = b * h * L / 1e9                       # [m³]
W_bt_kg    = V_bt_m3 * 2400                         # [kg]

# Cot chu duoi — 4phi18 (neo 2x250mm)
uw_main    = d_main**2 / 162                        # [kg/m]
L_bar_m    = (L + 2 * 250) / 1000                  # [m]
W_main_kg  = n_main * L_bar_m * uw_main

# Cot cau tao tren — 2phi12
uw_top     = d_top**2 / 162
W_top_kg   = n_top * L_bar_m * uw_top

# Cot dai — phi10@150
n_stir     = int((L - 100) / s_stir) + 1
inner_w    = b - 2 * a_cv                           # [mm]
inner_h    = h - 2 * a_cv                           # [mm]
L_1stir_m  = (2 * (inner_w + inner_h) + 2 * 10 * d_stir) / 1000
uw_stir    = d_stir**2 / 162
W_stir_kg  = n_stir * L_1stir_m * uw_stir

W_steel_kg = W_main_kg + W_top_kg + W_stir_kg

# Print ket qua
print("=" * 65)
print("  DAM BTCT B300x800 — L=6000mm — TCVN 5574:2018")
print("=" * 65)
print(f"\n[HINH HOC]")
print(f"  b x h = {int(b)} x {int(h)} mm  |  L = {int(L)} mm")
print(f"  a_s = {a_s:.1f} mm  |  h0 = {h0:.1f} mm")
print(f"\n[COT THEP]")
print(f"  Cot chu duoi : {n_main}phi{int(d_main)}  As = {As_main:.1f} mm2")
print(f"  Cot cau tao  : {n_top}phi{int(d_top)}    As_top = {As_top:.1f} mm2")
print(f"  As_min       : {As_min:.1f} mm2 -> {'OK' if ok_Asmin else 'FAIL'}")
print(f"  As yeu cau   : {As_req:.1f} mm2")
print(f"  Ham luong thep mu: {mu_pct:.3f} %")
print(f"\n[KIEM TOAN UON]")
print(f"  x = {x_act:.1f} mm  |  xi = {xi_act:.4f}  xi_R = {xi_R}  {'OK' if ok_xi else 'FAIL'}")
print(f"  MRd = {MRd:.2f} kN.m  |  Mu = {Mu:.1f} kN.m")
print(f"  Ti le Mu/MRd = {ratio_M:.3f}  ->  {'DAT ✅' if ok_bend else 'KHONG DAT ❌'}")
if not ok_bend:
    print(f"  GOI Y: {goi_y}")
print(f"\n[KHOI LUONG]")
print(f"  V be tong     : {V_bt_m3:.4f} m3  ({W_bt_kg:.0f} kg)")
print(f"  W cot chu     : {W_main_kg:.2f} kg  (n={n_main}, L={L_bar_m:.2f}m)")
print(f"  W cot cau tao : {W_top_kg:.2f} kg  (n={n_top})")
print(f"  W cot dai     : {W_stir_kg:.2f} kg  (n={n_stir} cai)")
print(f"  TONG THEP     : {W_steel_kg:.2f} kg")
print("=" * 65)

# ═══════════════════════════════════════════════════════════════
# 5. VE TRONG AUTOCAD (COM API)
# ═══════════════════════════════════════════════════════════════
try:
    import win32com.client, pythoncom
    from win32com.client import VARIANT

    def pt3d(x, y, z=0.0):
        return VARIANT(pythoncom.VT_ARRAY | pythoncom.VT_R8, (float(x), float(y), float(z)))

    def vtarr(*coords):
        flat = [float(v) for pair in coords for v in pair]
        return VARIANT(pythoncom.VT_ARRAY | pythoncom.VT_R8, flat)

    acad = win32com.client.GetActiveObject("AutoCAD.Application")
    # Tao ban ve moi neu chua co document nao active
    try:
        adoc = acad.ActiveDocument
        _    = adoc.Name   # kiem tra con song
    except Exception:
        import time
        adoc = acad.Documents.Add()
        time.sleep(3)
        adoc = acad.ActiveDocument
    msp      = adoc.ModelSpace
    out_dwg  = os.path.join(OUT_DIR, "Dam_B300x800_L6000.dwg")

    # --- Layers ---
    def mk_layer(name, color_idx, lw=0):
        try:
            lyr = adoc.Layers.Item(name)
        except:
            lyr = adoc.Layers.Add(name)
        lyr.Color = color_idx
        return lyr

    mk_layer("CONCRETE",  8)    # xam sang
    mk_layer("REBAR-MAIN", 1)   # do
    mk_layer("REBAR-TOP",  3)   # xanh la
    mk_layer("REBAR-STIR", 4)   # xanh duong
    mk_layer("DIM",        2)   # vang
    mk_layer("TEXT",       9)   # xam

    def set_lyr(name):
        adoc.ActiveLayer = adoc.Layers.Item(name)

    # ─── A. ELEVATION VIEW (6000 x 800mm) ───────────────────────
    EX, EY = 0.0, 1500.0   # origin mat dung (y=1500 co cho dim)

    set_lyr("CONCRETE")
    pl = msp.AddLightWeightPolyline(vtarr(
        (EX, EY), (EX+L, EY), (EX+L, EY+h), (EX, EY+h)
    ))
    pl.Closed = True; pl.Layer = "CONCRETE"

    # Cot chu duoi (line ngang)
    set_lyr("REBAR-MAIN")
    y_bot = EY + a_s
    ln1 = msp.AddLine(pt3d(EX+100, y_bot), pt3d(EX+L-100, y_bot))
    ln1.Layer = "REBAR-MAIN"

    # Cot cau tao tren (line ngang)
    set_lyr("REBAR-TOP")
    y_top_el = EY + h - (a_cv + d_stir + d_top/2)
    ln2 = msp.AddLine(pt3d(EX+100, y_top_el), pt3d(EX+L-100, y_top_el))
    ln2.Layer = "REBAR-TOP"

    # Cot dai (lines doc)
    set_lyr("REBAR-STIR")
    xs = EX + 50.0
    while xs <= EX + L - 50.0:
        ln_s = msp.AddLine(pt3d(xs, EY + a_cv), pt3d(xs, EY + h - a_cv))
        ln_s.Layer = "REBAR-STIR"
        xs += s_stir

    # ─── B. CROSS-SECTION 1-1 (300 x 800mm) ─────────────────────
    SX, SY = EX + L + 700.0, EY   # dat sang phai mat dung

    set_lyr("CONCRETE")
    pl_s = msp.AddLightWeightPolyline(vtarr(
        (SX, SY), (SX+b, SY), (SX+b, SY+h), (SX, SY+h)
    ))
    pl_s.Closed = True; pl_s.Layer = "CONCRETE"

    # Dai (hinh chu nhat trong)
    set_lyr("REBAR-STIR")
    pl_d = msp.AddLightWeightPolyline(vtarr(
        (SX+a_cv, SY+a_cv),
        (SX+b-a_cv, SY+a_cv),
        (SX+b-a_cv, SY+h-a_cv),
        (SX+a_cv, SY+h-a_cv)
    ))
    pl_d.Closed = True; pl_d.Layer = "REBAR-STIR"

    # Cot chu duoi — 4phi18 (hang duoi)
    set_lyr("REBAR-MAIN")
    x0_b = SX + a_cv + d_stir + d_main/2
    x1_b = SX + b - a_cv - d_stir - d_main/2
    sp_b = (x1_b - x0_b) / (n_main - 1) if n_main > 1 else 0
    y_b  = SY + a_s
    for i in range(n_main):
        cx = x0_b + i * sp_b
        c = msp.AddCircle(pt3d(cx, y_b), d_main/2)
        c.Layer = "REBAR-MAIN"

    # Cot cau tao tren — 2phi12
    set_lyr("REBAR-TOP")
    y_tp = SY + h - (a_cv + d_stir + d_top/2)
    for xi_i in [x0_b, x1_b]:
        c2 = msp.AddCircle(pt3d(xi_i, y_tp), d_top/2)
        c2.Layer = "REBAR-TOP"

    # ─── C. ANNOTATIONS ─────────────────────────────────────────
    TH = 90    # chieu cao chu [mm]
    TH_S = 70

    def txt(s, x, y, ht=None, layer="TEXT"):
        t = msp.AddText(str(s), pt3d(x, y), ht or TH)
        t.Layer = layer
        return t

    set_lyr("TEXT")
    # Tieu de mat dung
    txt("MAT DUNG DAM — b=300 x h=800mm, L=6000mm", EX, EY+h+250, TH)
    txt(f"Cot chu: {n_main}phi{int(d_main)} CB400-V  |  Cot dai: phi{int(d_stir)}@{int(s_stir)} CB240-T  |  Be tong: B25",
        EX, EY+h+140, TH_S)
    # Ket qua
    status_str = "DAT" if ok_bend else "KHONG DAT — AS YEU CAU = " + str(int(As_req)) + " mm2"
    txt(f"MRd = {MRd:.1f} kN.m  |  Mu = {Mu:.0f} kN.m  =>  {status_str}",
        EX, EY - 150, TH_S)
    txt(f"V be tong = {V_bt_m3:.4f} m3  |  Tong thep = {W_steel_kg:.1f} kg  (phi18={W_main_kg:.1f} + phi10dai={W_stir_kg:.1f} + phi12top={W_top_kg:.1f})",
        EX, EY - 280, TH_S)

    # Tieu de mat cat
    txt("MAT CAT 1-1   b=300  h=800 (mm)", SX, SY+h+250, TH)
    txt(f"{n_main}phi{int(d_main)} (duoi)  +  {n_top}phi{int(d_top)} (tren)  +  phi{int(d_stir)}@{int(s_stir)} (dai)",
        SX, SY+h+140, TH_S)

    # Ky hieu phi (text)
    txt(f"phi{int(d_main)}", x0_b - 50, y_b - 60, TH_S)
    txt(f"phi{int(d_top)}",  x0_b - 50, y_tp + 20, TH_S)
    txt(f"phi{int(d_stir)}@{int(s_stir)}", SX + b + 30, SY + h/2, TH_S)

    # ─── D. KICH THUOC (Dimensions) ─────────────────────────────
    set_lyr("DIM")

    def add_dim_linear(x1, y1, x2, y2, dim_x, dim_y, angle=0.0):
        try:
            d_obj = msp.AddDimLinear(
                pt3d(x1, y1), pt3d(x2, y2), pt3d(dim_x, dim_y), angle)
            d_obj.Layer = "DIM"
        except Exception as e_dim:
            pass  # khong loi neu dim loi

    # Chieu dai dam
    add_dim_linear(EX, EY, EX+L, EY, EX+L/2, EY-500)
    # Chieu cao dam (mat dung)
    add_dim_linear(EX, EY, EX, EY+h, EX-400, EY+h/2, 90.0)
    # Chieu rong mat cat
    add_dim_linear(SX, SY, SX+b, SY, SX+b/2, SY-500)
    # Chieu cao mat cat
    add_dim_linear(SX+b, SY, SX+b, SY+h, SX+b+400, SY+h/2, 90.0)

    # ─── E. ZOOM & SAVE ─────────────────────────────────────────
    adoc.SendCommand("ZOOM\nE\n")
    adoc.SendCommand("LAYERCLOSE\n")
    adoc.SaveAs(out_dwg)
    print(f"\n[DWG] Da luu: {out_dwg}")
    acad_ok = True

except Exception as e_acad:
    acad_ok = False
    out_dwg = "(Loi AutoCAD)"
    print(f"\n[WARN] AutoCAD: {e_acad}")

# ═══════════════════════════════════════════════════════════════
# 6. BAO CAO WORD (python-docx)
# ═══════════════════════════════════════════════════════════════
out_docx = os.path.join(OUT_DIR, "Report_Dam_B300x800_L6000.docx")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # --- Style ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_heading(doc, text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
        return p

    def set_cell(cell, text, bold=False, color=None, size=11, bg=None):
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if bg:
            set_cell_bg(cell, bg)

    def add_table_row(table, cells_data, bold_first=False, bg=None):
        row = table.add_row()
        for i, val in enumerate(cells_data):
            set_cell(row.cells[i], val,
                     bold=(bold_first and i == 0),
                     size=11, bg=bg)
        return row

    # ─── TIEU DE ────────────────────────────────────────────────
    title = doc.add_heading("THUYẾT MINH TÍNH TOÁN DẦM BÊ TÔNG CỐT THÉP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Theo TCVN 5574:2018 — Kết cấu Bê tông và BTCT")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph(f"Ngày lập: {NOW}  |  Phần mềm: Python 3.x + AutoCAD Civil 3D 2027")
    doc.add_paragraph()

    # ─── 1. THONG SO DAU VAO ───────────────────────────────────
    add_heading(doc, "1. Thông số Đầu vào", 1)
    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    for c, h_txt in zip(t1.rows[0].cells, ["Thông số", "Giá trị", "Đơn vị"]):
        set_cell(c, h_txt, bold=True, color=RGBColor(0xFF,0xFF,0xFF), bg="4472C4")

    rows_in = [
        ("Chiều rộng dầm b  (giả định)", f"{int(b)}", "mm"),
        ("Chiều cao dầm h", f"{int(h)}", "mm"),
        ("Chiều dài (nhịp) L", f"{int(L)}", "mm"),
        ("Lớp bảo vệ bê tông a_cv", f"{int(a_cv)}", "mm"),
        ("Mô men tính toán Mu", f"{Mu:.0f}", "kN.m"),
        ("Cốt thép chủ dưới", f"{n_main}ϕ{int(d_main)}", "CB400-V"),
        ("Cốt cấu tạo trên", f"{n_top}ϕ{int(d_top)}", "CB400-V"),
        ("Cốt đai", f"ϕ{int(d_stir)} @ {int(s_stir)}", "CB240-T, 2 nhánh"),
        ("Bê tông cấp độ bền", "B25", "—"),
        ("Rb (cường độ nén tính toán)", f"{Rb}", "MPa"),
        ("Rbt (cường độ kéo tính toán)", f"{Rbt}", "MPa"),
        ("Rs (thép chủ)", f"{Rs}", "MPa"),
        ("Rsw (thép đai)", f"{Rsw}", "MPa"),
        ("ξ_R (giới hạn, CB400-V/B25)", f"{xi_R}", "—"),
    ]
    for rdata in rows_in:
        add_table_row(t1, rdata, bold_first=True)

    doc.add_paragraph()

    # ─── 2. DAC TRUNG MAT CAT ──────────────────────────────────
    add_heading(doc, "2. Đặc trưng Mặt cắt", 1)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    for c, h_txt in zip(t2.rows[0].cells, ["Thông số", "Giá trị", "Đơn vị"]):
        set_cell(c, h_txt, bold=True, color=RGBColor(0xFF,0xFF,0xFF), bg="4472C4")

    rows_sect = [
        ("Diện tích mặt cắt A = b×h", f"{A_sect:.0f}", "mm²"),
        ("As chủ dưới (4ϕ18)", f"{As_main:.2f}", "mm²"),
        ("As cấu tạo trên (2ϕ12)", f"{As_top:.2f}", "mm²"),
        ("As đai / cắt ngang", f"{As_stir:.2f}", "mm²"),
        ("Khoảng cách TTT → mép dưới a_s", f"{a_s:.1f}", "mm"),
        ("Chiều cao làm việc h₀ = h − a_s", f"{h0:.1f}", "mm"),
        ("Hàm lượng thép μ = As/(b×h₀)", f"{mu_pct:.3f}", "%"),
        ("As_min (0.13%×b×h₀)", f"{As_min:.1f}", "mm²"),
    ]
    for rdata in rows_sect:
        add_table_row(t2, rdata, bold_first=True)

    doc.add_paragraph()

    # ─── 3. KHOI LUONG ─────────────────────────────────────────
    add_heading(doc, "3. Khối lượng Vật liệu", 1)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    for c, h_txt in zip(t3.rows[0].cells, ["Hạng mục", "Số lượng", "Khối lượng đơn vị", "Tổng KL"]):
        set_cell(c, h_txt, bold=True, color=RGBColor(0xFF,0xFF,0xFF), bg="4472C4")

    add_table_row(t3, ["Bê tông B25",
                        f"b×h×L = {b/1000:.2f}×{h/1000:.2f}×{L/1000:.2f} m",
                        "2400 kg/m³",
                        f"{V_bt_m3:.4f} m³  ({W_bt_kg:.0f} kg)"], bold_first=True)
    add_table_row(t3, [f"Cốt chủ dưới {n_main}ϕ{int(d_main)}",
                        f"n={n_main}, L_bar={L_bar_m:.2f} m/cây",
                        f"{uw_main:.4f} kg/m", f"{W_main_kg:.2f} kg"], bold_first=True)
    add_table_row(t3, [f"Cốt cấu tạo trên {n_top}ϕ{int(d_top)}",
                        f"n={n_top}, L_bar={L_bar_m:.2f} m/cây",
                        f"{uw_top:.4f} kg/m", f"{W_top_kg:.2f} kg"], bold_first=True)
    add_table_row(t3, [f"Cốt đai ϕ{int(d_stir)}@{int(s_stir)}",
                        f"n={n_stir} cái, L_đai={L_1stir_m:.3f} m/cái",
                        f"{uw_stir:.4f} kg/m", f"{W_stir_kg:.2f} kg"], bold_first=True)
    add_table_row(t3, ["TỔNG THÉP", "", "", f"{W_steel_kg:.2f} kg"],
                   bold_first=True, bg="D6E4F0")

    doc.add_paragraph()

    # ─── 4. KIEM TOAN ──────────────────────────────────────────
    add_heading(doc, "4. Kiểm toán Dầm theo TCVN 5574:2018", 1)

    # Cong thuc
    doc.add_paragraph("Điều kiện chịu uốn (mục 8.1.2): M ≤ M_Rd = R_b × b × x × (h₀ − x/2)")
    doc.add_paragraph(f"x = Rs × As / (Rb × b) = {Rs} × {As_main:.1f} / ({Rb} × {b:.0f}) = {x_act:.2f} mm")
    doc.add_paragraph(f"ξ = x/h₀ = {x_act:.2f}/{h0:.1f} = {xi_act:.4f}  ≤  ξ_R = {xi_R}  →  {'ĐẠT ✔' if ok_xi else 'KHÔNG ĐẠT ✖'}")
    doc.add_paragraph(f"M_Rd = {Rb} × {b:.0f} × {x_act:.2f} × ({h0:.1f} − {x_act/2:.2f}) = {MRd_Nmm/1e6:.2f} kN.m")

    # Bang ket qua
    t4 = doc.add_table(rows=1, cols=4)
    t4.style = 'Table Grid'
    for c, h_txt in zip(t4.rows[0].cells, ["Nội dung kiểm tra", "Yêu cầu", "Thực tế", "Kết luận"]):
        set_cell(c, h_txt, bold=True, color=RGBColor(0xFF,0xFF,0xFF), bg="4472C4")

    checks = [
        ("Hàm lượng thép tối thiểu",
         f"As ≥ As_min = {As_min:.0f} mm²",
         f"As = {As_main:.0f} mm²",
         "ĐẠT ✔" if ok_Asmin else "KHÔNG ĐẠT ✖"),
        ("Điều kiện vùng nén không vượt hạn",
         f"ξ ≤ ξ_R = {xi_R}",
         f"ξ = {xi_act:.4f}",
         "ĐẠT ✔" if ok_xi else "KHÔNG ĐẠT ✖"),
        ("Khả năng chịu mô men",
         f"M_Rd ≥ Mu = {Mu:.0f} kN.m",
         f"M_Rd = {MRd:.1f} kN.m",
         "ĐẠT ✔" if ok_bend else "KHÔNG ĐẠT ✖"),
    ]
    for chk in checks:
        row_c = t4.add_row()
        for i, val in enumerate(chk):
            if i == 3:
                if "ĐẠT" in val and "KHÔNG" not in val:
                    set_cell(row_c.cells[i], val, bold=True,
                             color=RGBColor(0x27, 0x6E, 0x27), bg="C6EFCE")
                else:
                    set_cell(row_c.cells[i], val, bold=True,
                             color=RGBColor(0x9C, 0x00, 0x06), bg="FFC7CE")
            else:
                set_cell(row_c.cells[i], val)

    doc.add_paragraph()

    # ─── 5. KET LUAN ───────────────────────────────────────────
    add_heading(doc, "5. Kết luận và Kiến nghị", 1)

    if ok_bend:
        kl = doc.add_paragraph(
            f"✅  Dầm b×h = {int(b)}×{int(h)} mm với cốt thép {n_main}ϕ{int(d_main)} "
            f"ĐẠT yêu cầu chịu lực với Mu = {Mu:.0f} kN.m "
            f"(M_Rd = {MRd:.1f} kN.m, tỷ lệ = {ratio_M:.3f}).")
    else:
        kl = doc.add_paragraph(
            f"❌  Dầm b×h = {int(b)}×{int(h)} mm với cốt thép {n_main}ϕ{int(d_main)} "
            f"KHÔNG ĐẠT yêu cầu chịu lực với Mu = {Mu:.0f} kN.m "
            f"(M_Rd = {MRd:.1f} kN.m, cần As_req ≥ {As_req:.0f} mm²).")
        doc.add_paragraph(f"🔧  Kiến nghị thay bằng: {goi_y}")

    doc.add_paragraph(
        "⚠️  Lưu ý: Kiểm tra cắt (shear), độ võng và nứt cần cung cấp thêm Vu và giá trị tải trọng đặc trưng.")

    doc.add_paragraph()
    doc.add_paragraph(
        f"Bản vẽ DWG: {os.path.basename(out_dwg) if acad_ok else out_dwg}")
    doc.add_paragraph(f"Xuất lúc: {NOW}")

    doc.save(out_docx)
    print(f"[DOCX] Da luu: {out_docx}")

except ImportError:
    print("[WARN] python-docx chua cai. Chay: pip install python-docx")
    # Fallback: xuat text report
    out_txt = os.path.join(OUT_DIR, "Report_Dam_B300x800_L6000.txt")
    with open(out_txt, 'w', encoding='utf-8') as f:
        f.write(f"THUYET MINH TINH TOAN DAM BTCT\n")
        f.write(f"TCVN 5574:2018 | Ngay: {NOW}\n")
        f.write("=" * 60 + "\n")
        f.write(f"b={int(b)}mm  h={int(h)}mm  L={int(L)}mm\n")
        f.write(f"Cot chu: {n_main}phi{int(d_main)}  As={As_main:.1f}mm2\n")
        f.write(f"h0={h0:.1f}mm  x={x_act:.1f}mm  xi={xi_act:.4f}\n")
        f.write(f"MRd={MRd:.2f} kN.m  Mu={Mu:.0f} kN.m\n")
        f.write(f"Ket qua: {'DAT' if ok_bend else 'KHONG DAT'}\n")
        f.write(f"As yeu cau: {As_req:.0f} mm2\n")
        if not ok_bend:
            f.write(f"Goi y: {goi_y}\n")
        f.write(f"\nKhoi luong:\n")
        f.write(f"  V be tong = {V_bt_m3:.4f} m3\n")
        f.write(f"  Tong thep = {W_steel_kg:.2f} kg\n")
    print(f"[TXT] Da luu fallback: {out_txt}")
    out_docx = out_txt

except Exception as e_docx:
    print(f"[WARN] Word: {e_docx}")

print(f"\n[HOAN THANH]")
print(f"  DWG  : {out_dwg}")
print(f"  DOCX : {out_docx}")
